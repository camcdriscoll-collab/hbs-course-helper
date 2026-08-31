#!/usr/bin/env python3
"""
canvas_refresh.py — Smart cheat sheet refresh

Modes:
  --daily   Sync files + refresh Notes for sessions in the next 2 calendar days
  --weekly  Full forward scan: sync all courses, fill missing Notes for next 6 weeks

Never looks back — sessions with due dates in the past are always skipped.

Scheduled via launchd:
  Daily  5pm  →  canvas_refresh.py --daily
  Sunday 8am  →  canvas_refresh.py --weekly
"""

import argparse
import base64
import hashlib
import json
import os
import re
import sys
from datetime import datetime, timezone, timedelta
from pathlib import Path
from urllib.error import HTTPError, URLError
from urllib.parse import urlencode
from urllib.request import HTTPRedirectHandler, Request, build_opener, urlopen
import canvas_organize
import canvas_readings
import weekly_overview
import calendar_sync

# ── Path resolution (tolerates folder renames/moves) ─────────────────────────
sys.path.insert(0, str(Path(__file__).parent))
import path_config
_paths = path_config.resolve()

DEST_ROOT   = _paths["coursework_root"]
PROMPT_FILE = _paths["master_prompt"]
ENV_FILE    = _paths["env_file"] or Path("/dev/null")
_COURSES    = _paths["courses"]   # abbrev → {canvas_id, folder_path, ...}

# Build flat dicts for callers that need them
COURSES = {a: d["canvas_id"] for a, d in _COURSES.items() if d["folder_path"]}

CANVAS_BASE = "https://hbs.instructure.com/api/v1"
BOSTON = timezone(timedelta(hours=-4))   # EDT — fall term Sep–Nov
MODEL  = "claude-sonnet-4-6"

READING_EXTS = {".pdf", ".docx", ".pptx", ".doc", ".ppt", ".txt"}
SLIDE_EXTS   = {".pptx", ".ppt"}   # always routed to General/Slides/

# PDF token budget: Claude processes each PDF page as an image (~200K tokens/MB).
# Budget 800K tokens for PDFs, leaving headroom for prompt + response.
PDF_TOKENS_PER_MB = 200_000
MAX_PDF_TOKEN_BUDGET = 800_000
PDF_PAGE_LIMIT = 50   # PDFs over this page count are skipped; delete Notes and re-run to override

# ── Env / config ──────────────────────────────────────────────────────────────

def load_env() -> dict:
    env = {}
    if ENV_FILE.exists():
        for line in ENV_FILE.read_text().splitlines():
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                k, _, v = line.partition("=")
                env[k.strip()] = v.strip().strip('"').strip("'")
    return env

ENV = load_env()

def cfg(key: str) -> str:
    return os.getenv(key, ENV.get(key, ""))

def require(key: str) -> str:
    v = cfg(key)
    if not v:
        sys.exit(f"\nMissing {key} — add to {ENV_FILE} or set as env var.\n")
    return v

# ── Canvas API ────────────────────────────────────────────────────────────────

class _StripAuthOnRedirect(HTTPRedirectHandler):
    def redirect_request(self, req, fp, code, msg, headers, newurl):
        new_req = super().redirect_request(req, fp, code, msg, headers, newurl)
        if new_req:
            from urllib.parse import urlparse
            if urlparse(newurl).netloc != urlparse(req.full_url).netloc:
                new_req.remove_header("Authorization")
        return new_req

_opener = build_opener(_StripAuthOnRedirect())

def canvas_get(path: str, params: dict | None = None) -> list | dict:
    token = require("CANVAS_API_TOKEN")
    url = f"{CANVAS_BASE}/{path.lstrip('/')}"
    if params:
        url += "?" + urlencode(params)
    results = []
    while url:
        req = Request(url, headers={"Authorization": f"Bearer {token}"})
        try:
            resp = _opener.open(req, timeout=30)
        except HTTPError as e:
            print(f"    HTTP {e.code}: {url}")
            return []
        data = json.loads(resp.read())
        if isinstance(data, list):
            results.extend(data)
        else:
            return data
        link = resp.headers.get("Link", "")
        url = None
        for part in link.split(","):
            if 'rel="next"' in part:
                m = re.search(r"<(.+?)>", part)
                if m:
                    url = m.group(1)
    return results

def canvas_download(url: str, dest: Path) -> bool:
    """Download a Canvas file URL; returns True if newly downloaded."""
    if dest.exists():
        return False
    dest.parent.mkdir(parents=True, exist_ok=True)
    token = require("CANVAS_API_TOKEN")
    req = Request(url, headers={"Authorization": f"Bearer {token}"})
    try:
        resp = _opener.open(req, timeout=60)
        dest.write_bytes(resp.read())
        return True
    except (HTTPError, URLError, OSError) as e:
        print(f"      ✗ {dest.name}: {e}")
        return False

# ── Utilities ─────────────────────────────────────────────────────────────────

def safe_name(s: str) -> str:
    s = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "-", s)
    return re.sub(r"-{2,}", "-", s).strip(". -")[:200]

def boston_date(iso: str) -> datetime:
    return datetime.fromisoformat(iso.replace("Z", "+00:00")).astimezone(BOSTON)

def yymmdd(dt: datetime) -> str:
    return dt.strftime("%y%m%d")

def pdf_page_count(path: Path) -> int:
    """Return page count of a PDF. Uses pypdf if available, falls back to regex."""
    try:
        from pypdf import PdfReader
        return len(PdfReader(str(path)).pages)
    except ImportError:
        pass
    except Exception:
        return 0
    # Regex fallback: /Count N appears in the Pages tree (largest value = total)
    try:
        counts = re.findall(rb'/Count\s+(\d+)', path.read_bytes())
        return max((int(c) for c in counts), default=0)
    except Exception:
        return 0


def strip_html(html: str) -> str:
    text = re.sub(r"<[^>]+>", " ", html)
    for ent, rep in [("&nbsp;", " "), ("&amp;", "&"), ("&lt;", "<"), ("&gt;", ">")]:
        text = text.replace(ent, rep)
    return re.sub(r"[ \t]+", " ", text).strip()

def class_number(text: str) -> int | None:
    m = re.search(r"\bclass\s+(\d+)\b", text, re.IGNORECASE)
    return int(m.group(1)) if m else None

# ── Session discovery ─────────────────────────────────────────────────────────

def get_upcoming_sessions(horizon_days: int) -> list[dict]:
    """
    Return all assignments across all courses with due dates
    between now and now+horizon_days, sorted by due date.
    Each entry: {abbrev, course_id, assignment, date_str, due_dt}
    """
    now = datetime.now(tz=BOSTON)
    cutoff = now + timedelta(days=horizon_days)
    sessions = []

    for abbrev, course_id in COURSES.items():
        assignments = canvas_get(f"courses/{course_id}/assignments", {"per_page": 100})
        for a in assignments:
            if not a.get("due_at"):
                continue
            dt = boston_date(a["due_at"])
            if dt < now or dt > cutoff:
                continue
            sessions.append({
                "abbrev":    abbrev,
                "course_id": course_id,
                "assignment": a,
                "date_str":  yymmdd(dt),
                "due_dt":    dt,
            })

    sessions.sort(key=lambda s: s["due_dt"])
    return sessions

# ── File sync (targeted) ──────────────────────────────────────────────────────

def sync_course_files(course_id: int, abbrev: str, target_date_str: str | None = None):
    """
    Download Canvas files for a course.
    If target_date_str given, only sync files relevant to that session
    (files linked in the assignment + per-class Canvas folders).
    Otherwise sync all course files to General/.
    """
    course_folder = (_COURSES.get(abbrev, {}).get("folder_path") or DEST_ROOT / abbrev)
    general_dir = course_folder / "General"
    general_dir.mkdir(parents=True, exist_ok=True)

    folders = canvas_get(f"courses/{course_id}/folders", {"per_page": 100})
    class_folders: dict[int, dict] = {}
    class_folder_ids: set[int] = set()
    for f in folders:
        n = class_number(f.get("name", ""))
        if n is not None:
            class_folders[n] = f
            class_folder_ids.add(f["id"])

    all_files = canvas_get(f"courses/{course_id}/files", {"per_page": 100})
    file_by_id = {f["id"]: f for f in all_files}

    slides_dir = general_dir / "Slides"

    # Build index of filenames already placed on disk (session folders + Slides/ + Supplemental/)
    # to avoid re-downloading files that canvas_organize already placed correctly.
    _placed: set[str] = set()
    if course_folder.exists():
        for d in course_folder.iterdir():
            if d.is_dir() and re.match(r'^\d{6}\s', d.name):
                _placed.update(f2.name for f2 in d.iterdir() if f2.is_file())
        for subdir_name in ("Slides", "Supplemental"):
            subdir = general_dir / subdir_name
            if subdir.exists():
                _placed.update(f2.name for f2 in subdir.iterdir() if f2.is_file())

    # General files: route PPTX → General/Slides/, others → General/ root
    # Skip anything already placed on disk.
    for f in all_files:
        if f.get("folder_id") not in class_folder_ids:
            fname = safe_name(f["display_name"])
            if fname in _placed:
                continue  # already on disk in the right place
            if Path(fname).suffix.lower() in SLIDE_EXTS:
                slides_dir.mkdir(parents=True, exist_ok=True)
                dest = slides_dir / fname
            else:
                dest = general_dir / fname
            if canvas_download(f["url"], dest):
                print(f"    ↓ [General] {f['display_name']}")

    if not target_date_str:
        return

    # Targeted: sync files for the specific session
    assignments = canvas_get(f"courses/{course_id}/assignments", {"per_page": 100})
    for a in assignments:
        if not a.get("due_at"):
            continue
        if yymmdd(boston_date(a["due_at"])) != target_date_str:
            continue

        session_dir = course_folder / f"{target_date_str} {abbrev}"
        session_dir.mkdir(parents=True, exist_ok=True)
        cn = class_number(a.get("name", ""))

        # Per-class Canvas folder
        if cn and cn in class_folders:
            folder = class_folders[cn]
            for f in canvas_get(f"folders/{folder['id']}/files", {"per_page": 100}):
                fname = safe_name(f["display_name"])
                # PPTX always goes to General/Slides/, never in session folder
                if Path(fname).suffix.lower() in SLIDE_EXTS:
                    slides_dir.mkdir(parents=True, exist_ok=True)
                    dest = slides_dir / fname
                else:
                    dest = session_dir / fname
                if canvas_download(f["url"], dest):
                    print(f"    ↓ [canvas folder] {f['display_name']}")

        # Files linked in assignment description
        desc = a.get("description") or ""
        ids = re.findall(
            rf"instructure\.com/(?:courses/{course_id}/)?files/(\d+)", desc
        )
        for fid_str in ids:
            fid = int(fid_str)
            if fid in file_by_id:
                f = file_by_id[fid]
                fname = safe_name(f["display_name"])
                if Path(fname).suffix.lower() in SLIDE_EXTS:
                    slides_dir.mkdir(parents=True, exist_ok=True)
                    dest = slides_dir / fname
                else:
                    dest = session_dir / fname
                if canvas_download(f["url"], dest):
                    print(f"    ↓ [linked] {f['display_name']}")

# ── Markdown → docx conversion ───────────────────────────────────────────────

def _parse_inline(para, text: str) -> None:
    """Add runs to a paragraph with **bold** and *italic* applied. Always 12pt."""
    from docx.shared import Pt
    for seg in re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text):
        if seg.startswith('**') and seg.endswith('**'):
            r = para.add_run(seg[2:-2])
            r.bold = True
            r.font.size = Pt(12)
        elif seg.startswith('*') and seg.endswith('*'):
            r = para.add_run(seg[1:-1])
            r.italic = True
            r.font.size = Pt(12)
        elif seg:
            para.add_run(seg).font.size = Pt(12)


def _tight(para) -> None:
    """Enforce 12pt font and tight spacing on a paragraph."""
    from docx.shared import Pt
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(6)


def markdown_to_docx(md_text: str, output_path: Path,
                     title: str, metadata: dict) -> None:
    """Convert Claude's markdown output → formatted .docx file."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        sys.exit("python-docx not installed — run: uv pip install python-docx")

    doc = Document()

    # Force 12pt + tight spacing on all built-in styles up front
    for style_name in ('Normal', 'Title', 'Heading 1', 'Heading 2', 'Heading 3',
                       'List Bullet', 'List Bullet 2', 'List Number'):
        try:
            st = doc.styles[style_name]
            st.font.size = Pt(12)
            st.paragraph_format.space_before = Pt(0)
            st.paragraph_format.space_after  = Pt(6)
        except KeyError:
            pass

    # Margins: 1.25 in sides, 1 in top/bottom
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # Title
    p = doc.add_heading(title, level=0)
    _tight(p)

    # Metadata block
    for key, val in metadata.items():
        p = doc.add_paragraph()
        r = p.add_run(f"{key}: ")
        r.bold = True
        r.font.size = Pt(12)
        p.add_run(val).font.size = Pt(12)
        _tight(p)

    for line in md_text.split('\n'):
        s = line.rstrip()

        if not s:
            continue
        elif s.startswith('### '):
            p = doc.add_heading(level=3)
            _parse_inline(p, s[4:])
            _tight(p)
        elif s.startswith('## '):
            p = doc.add_heading(level=2)
            _parse_inline(p, s[3:])
            _tight(p)
        elif s.startswith('# '):
            p = doc.add_heading(level=1)
            _parse_inline(p, s[2:])
            _tight(p)
        elif re.match(r'^[-*_]{3,}$', s):
            # Horizontal rule via paragraph bottom border
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'AAAAAA')
            pBdr.append(bottom)
            pPr.append(pBdr)
            _tight(p)
        elif re.match(r'^    [-*] |^  [-*] ', line):
            p = doc.add_paragraph(style='List Bullet 2')
            _parse_inline(p, s.lstrip('*- \t'))
            _tight(p)
        elif s.startswith('- ') or s.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            _parse_inline(p, s[2:])
            _tight(p)
        elif re.match(r'^\d+\. ', s):
            p = doc.add_paragraph(style='List Number')
            _parse_inline(p, re.sub(r'^\d+\. ', '', s))
            _tight(p)
        else:
            p = doc.add_paragraph(style='Normal')
            _parse_inline(p, s)
            _tight(p)

    doc.save(str(output_path))


# ── Staleness check ───────────────────────────────────────────────────────────

def notes_are_stale(session_dir: Path, abbrev: str, date_str: str,
                    weekly: bool = False,
                    canvas_desc_hash: str = "",
                    skip_prompt_regen: bool = False) -> tuple[bool, str]:
    """
    Returns (should_regenerate, reason).
    Checks: missing Notes file, Canvas description changed, new reading files,
    prompt file updated since last generation.
    """
    # Check .docx first (new format), fall back to .md (old format)
    notes_file = session_dir / f"{date_str} {abbrev} Notes.docx"
    if not notes_file.exists():
        notes_file = session_dir / f"{date_str} {abbrev} Notes.md"
    if not notes_file.exists():
        return True, "no Notes file yet"

    notes_mtime = notes_file.stat().st_mtime

    # Canvas assignment description changed (professor edited the posting)
    if canvas_desc_hash:
        meta_file = session_dir / ".notes_meta.json"
        if meta_file.exists():
            try:
                stored = json.loads(meta_file.read_text())
                if stored.get("canvas_hash", "") != canvas_desc_hash:
                    return True, "Canvas assignment description changed"
            except Exception:
                pass

    # Any reading file newer than Notes file → stale
    reading_files = [
        f for f in session_dir.iterdir()
        if f.is_file()
        and f.suffix.lower() in READING_EXTS
        and "Notes" not in f.name
    ]
    for f in reading_files:
        if f.stat().st_mtime > notes_mtime:
            return True, f"new reading: {f.name}"

    # Master prompt or per-course refinement prompt updated → stale
    # (only for future sessions — past sessions are excluded by the caller)
    if not skip_prompt_regen:
        code = abbrev.replace(" ", "_")
        refinement = (_COURSES.get(abbrev, {}).get("refinement_prompt")
                      or path_config.PROMPTS_DIR / f"cheat_sheet_prompt_{code}_refinement.md")
        for prompt_f in [PROMPT_FILE, refinement]:
            if prompt_f and prompt_f.exists():
                if prompt_f.stat().st_mtime > notes_mtime:
                    return True, f"prompt updated: {prompt_f.name}"

    return False, "up to date"

# ── Cheat sheet generation ────────────────────────────────────────────────────

def generate_notes(session: dict):
    abbrev    = session["abbrev"]
    date_str  = session["date_str"]
    assignment = session["assignment"]
    course_folder = (_COURSES.get(abbrev, {}).get("folder_path") or DEST_ROOT / abbrev)
    session_dir = course_folder / f"{date_str} {abbrev}"
    output_file = session_dir / f"{date_str} {abbrev} Notes.docx"

    session_dir.mkdir(parents=True, exist_ok=True)

    # Sort: largest PDFs first (proxy for "main case"), non-PDFs after
    reading_files = sorted(
        (f for f in session_dir.iterdir()
         if f.is_file() and f.suffix.lower() in READING_EXTS and "Notes" not in f.name),
        key=lambda f: (-f.stat().st_size if f.suffix.lower() == ".pdf" else 0, f.name),
    )

    # Build prompt
    master = PROMPT_FILE.read_text() if PROMPT_FILE.exists() else ""
    code = abbrev.replace(" ", "_")
    notes_file = (_COURSES.get(abbrev, {}).get("refinement_prompt")
                  or path_config.PROMPTS_DIR / f"cheat_sheet_prompt_{code}_refinement.md")
    if notes_file.exists():
        raw = re.sub(r"<!--.*?-->", "", notes_file.read_text(), flags=re.DOTALL).strip()
        if raw and raw != "# CLASS-SPECIFIC NOTES":
            master = re.sub(
                r"\[CLASS-SPECIFIC NOTES.*?\].*",
                f"[CLASS-SPECIFIC NOTES]\n{raw}",
                master, flags=re.DOTALL,
            )

    canvas_block = ""
    if assignment:
        name = assignment.get("name", "")
        desc = strip_html(assignment.get("description") or "")
        canvas_block = f"\n\n=== CANVAS ASSIGNMENT POSTING ===\nTitle: {name}\n\n{desc}"

    prompt_text = master + canvas_block

    if not reading_files and not assignment:
        print(f"    ⚠ Nothing to generate for {date_str} {abbrev} — skipping")
        return

    # Build message content
    try:
        import anthropic as ant
    except ImportError:
        sys.exit("anthropic not installed — run: pip install anthropic")

    api_key = require("ANTHROPIC_API_KEY")
    client  = ant.Anthropic(api_key=api_key)

    content: list[dict] = []
    skipped: list[str] = []
    pdf_token_used = 0
    if reading_files:
        content.append({"type": "text", "text": "Here are the assigned readings:"})
        for f in reading_files:
            if f.suffix.lower() == ".pdf":
                pages = pdf_page_count(f)
                if pages > PDF_PAGE_LIMIT:
                    print(f"    ⚠ Skipped ({pages}p > {PDF_PAGE_LIMIT}-page limit): {f.name}")
                    skipped.append(f"{f.name} ({pages}p, too long)")
                    content.append({"type": "text", "text": (
                        f"=== {f.name} ===\n"
                        f"[Skipped: {pages} pages exceeds the {PDF_PAGE_LIMIT}-page limit. "
                        f"Delete the Notes file and re-run to force inclusion.]"
                    )})
                    continue
                size_mb = f.stat().st_size / (1024 * 1024)
                estimated_tokens = int(size_mb * PDF_TOKENS_PER_MB)
                if pdf_token_used + estimated_tokens > MAX_PDF_TOKEN_BUDGET:
                    print(f"    ⚠ Skipped (token budget, ~{estimated_tokens//1000}k tokens): {f.name}")
                    skipped.append(f.name)
                    content.append({"type": "text", "text": (
                        f"=== {f.name} ===\n"
                        f"[File omitted to stay within context limit (~{size_mb:.1f} MB / "
                        f"~{estimated_tokens//1000}k tokens). Summarize from Canvas description.]"
                    )})
                    continue
                pdf_token_used += estimated_tokens
                data = base64.standard_b64encode(f.read_bytes()).decode()
                content.append({
                    "type": "document",
                    "source": {"type": "base64", "media_type": "application/pdf", "data": data},
                    "title": f.name,
                })
            elif f.suffix.lower() in {".pptx", ".ppt"}:
                # Binary slide files — mention but don't read raw bytes
                print(f"    – Slides (listed but not read as text): {f.name}")
                skipped.append(f.name)
                content.append({"type": "text", "text": (
                    f"=== {f.name} ===\n"
                    "[Slide deck — not included as raw text. "
                    "Summarize from Canvas description and any PDFs.]"
                )})
            else:
                # Text-based files (.txt, .md, .docx, etc.) — read but cap size
                try:
                    text = f.read_text(errors="replace")
                except Exception:
                    text = f"[Could not read {f.name}]"
                if len(text) > 400_000:  # ~100k tokens
                    text = text[:400_000] + "\n[... truncated — file too large ...]"
                content.append({"type": "text", "text": f"=== {f.name} ===\n{text}"})

    content.append({"type": "text", "text": prompt_text})

    msg = client.messages.create(
        model=MODEL, max_tokens=8192,
        messages=[{"role": "user", "content": content}],
    )
    if msg.stop_reason == "max_tokens":
        print(f"    ⚠ Output truncated (hit max_tokens limit) — consider splitting readings")
    cost = (msg.usage.input_tokens * 3 + msg.usage.output_tokens * 15) / 1_000_000
    print(f"    Tokens: {msg.usage.input_tokens:,} in / {msg.usage.output_tokens:,} out  (~${cost:.3f})")
    result = msg.content[0].text

    # Save canvas hash for staleness detection on future runs
    canvas_hash = ""
    if assignment:
        canvas_hash = hashlib.md5(
            strip_html(assignment.get("description") or "").encode()
        ).hexdigest()[:12]
    (session_dir / ".notes_meta.json").write_text(json.dumps({"canvas_hash": canvas_hash}))

    metadata: dict[str, str] = {
        "Generated": datetime.now().strftime("%Y-%m-%d %H:%M"),
    }
    if assignment:
        metadata["Canvas"] = assignment["name"]
    if reading_files:
        included = [f.name for f in reading_files if f.name not in skipped]
        if included:
            metadata["Readings"] = ", ".join(included)
    if skipped:
        metadata["Skipped"] = ", ".join(skipped)

    markdown_to_docx(
        md_text     = result,
        output_path = output_file,
        title       = f"{date_str} {abbrev} Notes",
        metadata    = metadata,
    )
    print(f"    ✅ Generated: {output_file.name}")

# ── Podcast generation (wraps podcast_gen.py) ─────────────────────────────────

def generate_podcast_for_session(session: dict):
    """
    Generate a NotebookLM podcast for a session (synchronous wrapper).
    Skipped if the .m4a already exists.
    Requires notebooklm-py and a valid ~/.notebooklm session.
    """
    import asyncio
    abbrev    = session["abbrev"]
    date_str  = session["date_str"]
    course_folder = (_COURSES.get(abbrev, {}).get("folder_path") or DEST_ROOT / abbrev)
    session_dir   = course_folder / f"{date_str} {abbrev}"
    podcast_file  = session_dir / f"{date_str} {abbrev} Podcast.m4a"

    if podcast_file.exists():
        print(f"    ✓ Podcast exists: {podcast_file.name}")
        return

    try:
        import podcast_gen as _pg
    except ImportError:
        print("    ⚠ podcast_gen.py not found on sys.path — skipping podcast")
        return
    try:
        asyncio.run(_pg._generate(date_str, abbrev))
    except Exception as e:
        print(f"    ✗ Podcast generation failed: {e}")


# ── Modes ─────────────────────────────────────────────────────────────────────

def run_daily(skip_prompt_regen: bool = False, with_podcast: bool = False):
    """Sync files + refresh Notes for sessions in the next 2 calendar days."""
    now = datetime.now(tz=BOSTON)
    today = now.date()
    # "next 2 calendar days" = today and tomorrow
    cutoff_date = today + timedelta(days=2)

    print(f"\n{'─'*55}")
    print(f"  DAILY REFRESH — sessions through {cutoff_date}")
    print(f"{'─'*55}")

    sessions = get_upcoming_sessions(horizon_days=2)

    if not sessions:
        print("  No upcoming sessions in the next 2 days.")
        return

    for s in sessions:
        abbrev   = s["abbrev"]
        date_str = s["date_str"]
        label    = s["assignment"].get("name", f"{date_str} {abbrev}")[:60]
        print(f"\n  [{date_str}] {abbrev} — {label}")

        # Sync Canvas files + externally-linked readings for this session
        sync_course_files(s["course_id"], abbrev, target_date_str=date_str)

        course_folder = (_COURSES.get(abbrev, {}).get("folder_path") or DEST_ROOT / abbrev)
        session_dir = course_folder / f"{date_str} {abbrev}"
        n_read = canvas_readings.sync_reading_links(s["assignment"], session_dir)
        if n_read:
            print(f"    ↓ {n_read} reading(s) saved")

        desc_text = strip_html(s["assignment"].get("description") or "")
        canvas_hash = hashlib.md5(desc_text.encode()).hexdigest()[:12]
        stale, reason = notes_are_stale(
            session_dir, abbrev, date_str, weekly=False,
            canvas_desc_hash=canvas_hash, skip_prompt_regen=skip_prompt_regen,
        )
        if stale:
            print(f"    → Regenerating Notes ({reason})")
            generate_notes(s)
        else:
            print(f"    ✓ Notes up to date")

        if with_podcast:
            print(f"    Podcast check...")
            generate_podcast_for_session(s)

    print("\n  Organizing folders...")
    canvas_organize.organize_all(verbose=True)

    print("\n  Checking for duplicates...")
    trashed = canvas_organize.dedup_to_trash(verbose=True)
    if trashed:
        print(f"  {trashed} duplicate(s) moved to Trash.")

    print("\n  Syncing calendar...")
    calendar_sync.run()

    print(f"\n{'─'*55}")
    print("  Daily refresh complete.")
    print(f"{'─'*55}\n")


def run_weekly(skip_prompt_regen: bool = False, with_podcast: bool = False):
    """
    Full forward scan:
      - Sync files for all courses (6-week horizon)
      - Generate Notes only for sessions within the next 2 weeks
        (avoids burning compute on sessions that may still change)
    """
    now = datetime.now(tz=BOSTON)
    notes_cutoff = now + timedelta(days=14)

    print(f"\n{'─'*55}")
    print(f"  WEEKLY REFRESH — sync 6 weeks, notes ≤ 2 weeks")
    print(f"{'─'*55}")

    # Full file sync for all courses (6-week horizon)
    print("\n  Syncing all course files...")
    for abbrev, course_id in COURSES.items():
        print(f"  {abbrev}...")
        sync_course_files(course_id, abbrev, target_date_str=None)

    # Notes: only sessions in the next 2 weeks
    all_sessions = get_upcoming_sessions(horizon_days=42)
    notes_sessions = [s for s in all_sessions if s["due_dt"] <= notes_cutoff]
    later_sessions = [s for s in all_sessions if s["due_dt"] > notes_cutoff]

    print(f"\n  Upcoming sessions: {len(all_sessions)} total, "
          f"{len(notes_sessions)} within 2 weeks (notes), "
          f"{len(later_sessions)} later (files only).")

    for s in all_sessions:
        abbrev   = s["abbrev"]
        date_str = s["date_str"]
        label    = s["assignment"].get("name", f"{date_str} {abbrev}")[:60]
        course_folder = (_COURSES.get(abbrev, {}).get("folder_path") or DEST_ROOT / abbrev)
        session_dir = course_folder / f"{date_str} {abbrev}"

        if s["due_dt"] > notes_cutoff:
            print(f"  [{date_str}] {abbrev} — files only (>2 weeks out)")
            continue

        # Always sync Canvas folder files + external reading links for sessions in window
        print(f"\n  [{date_str}] {abbrev} — {label}")
        sync_course_files(s["course_id"], abbrev, target_date_str=date_str)
        n_read = canvas_readings.sync_reading_links(s["assignment"], session_dir)
        if n_read:
            print(f"    ↓ {n_read} reading(s) saved")

        desc_text = strip_html(s["assignment"].get("description") or "")
        canvas_hash = hashlib.md5(desc_text.encode()).hexdigest()[:12]
        stale, reason = notes_are_stale(
            session_dir, abbrev, date_str, weekly=True,
            canvas_desc_hash=canvas_hash, skip_prompt_regen=skip_prompt_regen,
        )
        if stale:
            print(f"    → Regenerating Notes ({reason})")
            generate_notes(s)
        else:
            print(f"    ✓ Notes up to date")

    print("\n  Organizing folders...")
    canvas_organize.organize_all(verbose=True)

    print("\n  Checking for duplicates...")
    trashed = canvas_organize.dedup_to_trash(verbose=True)
    if trashed:
        print(f"  {trashed} duplicate(s) moved to Trash.")
    else:
        print("  No duplicates found.")

    print("\n  Generating weekly overview...")
    ov = weekly_overview.generate()
    print(f"  Saved: {ov}")

    print("\n  Syncing calendar...")
    calendar_sync.run()

    if with_podcast:
        import subprocess
        subprocess.run(["open", str(ov)], check=False)

        print(f"\n{'─'*55}")
        print(f"  PODCAST GENERATION — {len(notes_sessions)} session(s) in range")
        print(f"{'─'*55}")
        for i, s in enumerate(notes_sessions, 1):
            day_label = s["due_dt"].strftime("%a %b %-d")
            name  = s["assignment"].get("name", f"{s['date_str']} {s['abbrev']}")
            parts = name.split("|")
            short = (parts[-1].strip() if len(parts) > 1 else name.strip())[:55]
            print(f"  {i:2}. {day_label}  {s['abbrev']:<10} — {short}")

        print(f"\n  Enter numbers to SKIP (comma-separated), or press Enter for all:")
        try:
            raw = input("  > ").strip()
        except (EOFError, KeyboardInterrupt):
            raw = ""

        skip_indices: set[int] = set()
        for part in raw.split(","):
            part = part.strip()
            if part.isdigit():
                idx = int(part) - 1
                if 0 <= idx < len(notes_sessions):
                    skip_indices.add(idx)

        print()
        for i, s in enumerate(notes_sessions):
            if i in skip_indices:
                print(f"  – Skipped: {s['abbrev']} {s['date_str']}")
            else:
                generate_podcast_for_session(s)

    print(f"\n{'─'*55}")
    print("  Weekly refresh complete.")
    print(f"{'─'*55}\n")

# ── Entry point ───────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description=__doc__)
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--daily",  action="store_true", help="Next 2 calendar days")
    group.add_argument("--weekly", action="store_true", help="Full 6-week forward scan")
    parser.add_argument("--skip-prompt-regen", action="store_true",
                        help="Don't regenerate notes just because the master prompt changed "
                             "(useful after minor prompt tweaks)")
    parser.add_argument("--with-podcast", action="store_true",
                        help="Also generate NotebookLM podcasts for sessions within the notes "
                             "window (requires notebooklm login; adds ~10 min per session)")
    args = parser.parse_args()

    if args.daily:
        run_daily(skip_prompt_regen=args.skip_prompt_regen, with_podcast=args.with_podcast)
    else:
        run_weekly(skip_prompt_regen=args.skip_prompt_regen, with_podcast=args.with_podcast)


if __name__ == "__main__":
    main()
